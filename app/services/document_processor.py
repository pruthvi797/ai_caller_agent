"""
Document Processing Service — Universal Multimodal Extraction
=============================================================

Handles ALL car dealership document types reliably:
  - Car Brochures         (PDF, DOCX, images) — multi-page, image-heavy, variant tables
  - Pricing Sheets        (PDF, DOCX)         — variant × price grids, ex-showroom/on-road
  - Feature Comparisons   (PDF, DOCX)         — tick/cross grids across variants
  - Promotional Offers    (PDF, images)       — discounts, EMI schemes, exchange bonuses
  - Spec Sheets           (PDF, DOCX)         — engine, dimensions, transmission data
  - Warranty / Service    (PDF, DOCX)         — warranty terms, service schedule, AMC
  - Test Drive Reports    (PDF, DOCX)         — drive experience, feedback summaries
  - Insurance Documents   (PDF)              — coverage details, premium breakdowns
  - Accessory Catalogues  (PDF, images)      — accessory names, prices, fitment info
  - General / Mixed       (any)              — fallback for unknown types

Pipeline:
  PDF   → render each page to PNG → OpenAI Vision (type-specific prompt)
            → If OpenAI fails → PyMuPDF text fallback
  DOCX  → python-docx text + tables; embedded images → OpenAI Vision
  Image → OpenAI Vision directly (or warn if PYMUPDF_ONLY mode)

  In PYMUPDF_ONLY mode (dev/testing):
    - PDFs  → PyMuPDF raw text (fast, free, no API calls)
    - DOCX  → python-docx text only (no embedded image extraction)
    - Images → Warning logged (Vision API required for images)

Robustness features:
  - Auto document-type detection from upload field + filename + content signals
  - Boilerplate / legal disclaimer filtering (removes noise chunks)
  - Smart chunking: section-aware splits, no orphan fragments
  - Section classification: pricing / features / safety / specs / colors / offers / warranty
  - processed_at timestamp correctly set on completion
  - Full error recovery with DB rollback on failure

Install:
  pip install openai pymupdf pillow python-docx --break-system-packages

.env:
  OPENAI_API_KEY=sk-...
  PYMUPDF_ONLY=true     # set false for production (enables Vision API)
"""

import os
import re
import time
import uuid
import base64
import hashlib
import logging
from typing import List, Tuple, Optional, Dict
from datetime import datetime
from sqlalchemy.orm import Session

from app.models.document_model import Document
from app.models.document_chunk_model import DocumentChunk

logger = logging.getLogger(__name__)

# Load .env before reading os.getenv() — background tasks import this module
# before database.py runs load_dotenv(), so we must call it ourselves.
try:
    from dotenv import load_dotenv as _load_dotenv
    _load_dotenv()
except ImportError:
    pass


# ══════════════════════════════════════════════════════════════════════════════
# CONFIG
# ══════════════════════════════════════════════════════════════════════════════

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# PYMUPDF_ONLY=true  → free/fast dev mode (no OpenAI API calls)
# PYMUPDF_ONLY=false → production mode (OpenAI Vision primary)
PYMUPDF_ONLY: bool = os.getenv("PYMUPDF_ONLY", "true").lower() in ("1", "true", "yes")

# Model:
#   "gpt-4o-mini" → cheapest, fast, good for most docs    ~$0.00015/image
#   "gpt-4o"      → best quality for dense tables/grids   ~$0.00150/image
OPENAI_MODEL = "gpt-4o-mini"

PAGE_DPI           = 150    # PDF render resolution — 150 is good balance of quality vs size
MAX_PAGES          = 40     # Safety cap for very large documents
CHUNK_SIZE         = 800    # Target chars per chunk (soft limit)
MIN_CHUNK_CHARS    = 120    # Minimum chars — filters out fragment/orphan chunks
MIN_TEXT_THRESHOLD = 10     # Min chars for PyMuPDF page text to be considered non-empty

OPENAI_MIN_INTERVAL_SECONDS = 0.5  # Rate-limit guard between API calls

# In-process page cache: MD5(image_bytes) → extracted text
# Prevents duplicate API calls on reprocess of same document
_page_cache: Dict[str, str] = {}

# Quota guard: flip to True on first quota error; all pages fall back to PyMuPDF
_openai_quota_exhausted: bool = False
_last_openai_call_time: float = 0.0


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT TYPE DETECTION
# ══════════════════════════════════════════════════════════════════════════════

# Keyword signals used to auto-detect document type from extracted text
# (used as fallback when upload field = "other" or filename is ambiguous)
DOC_TYPE_CONTENT_SIGNALS = {
    "pricing_sheet": [
        "ex-showroom", "on-road price", "variant price", "price list",
        "₹", "lakh", "booking amount", "rto charges", "insurance charges",
        "extended warranty cost", "accessories cost",
    ],
    "feature_comparison": [
        "features list", "feature comparison", "available", "not available",
        "lxi mt", "vxi mt", "zxi mt", "zxi+", "alpha", "delta", "sigma", "zeta",
        "legend:", "tick", "variant wise features",
    ],
    "promotional_offer": [
        "cash discount", "exchange bonus", "loyalty bonus", "corporate discount",
        "cashback", "festive offer", "limited period", "save up to", "emi scheme",
        "zero down payment", "subvention", "benefit of",
    ],
    "spec_sheet": [
        "wheelbase", "kerb weight", "ground clearance", "fuel tank capacity",
        "displacement", "max power", "max torque", "emission norm", "bs vi",
        "ventilated disc", "torsion beam", "macpherson", "tyre size",
    ],
    "warranty": [
        "warranty period", "extended warranty", "free service", "service schedule",
        "annual maintenance", "amc", "roadside assistance", "coverage",
        "whichever is earlier", "labour charges",
    ],
    "insurance": [
        "sum insured", "idv", "insured declared value", "premium", "ncb",
        "no claim bonus", "third party", "comprehensive", "zero depreciation",
        "claim process", "policy number", "insurer",
    ],
    "accessory": [
        "accessory", "accessories", "fitment", "genuine parts", "add-on",
        "floor mat", "seat cover", "body cover", "dash cam", "alloy wheel price",
    ],
}

# Filename keywords → document type
FILENAME_TYPE_MAP = {
    "pricing_sheet":      ["price", "pricing", "rate", "cost", "tariff"],
    "feature_comparison": ["feature", "comparison", "compare", "spec", "specification"],
    "promotional_offer":  ["offer", "discount", "promo", "scheme", "festive", "cashback"],
    "warranty":           ["warranty", "service", "amc", "maintenance"],
    "insurance":          ["insurance", "insur", "policy"],
    "accessory":          ["accessory", "accessories", "addon", "add-on", "parts"],
    "brochure":           ["brochure", "catalogue", "catalog", "booklet"],
}

# Upload field → extraction prompt key
UPLOAD_FIELD_MAP = {
    "brochure":           "brochure",
    "pricing_sheet":      "pricing_sheet",
    "feature_comparison": "feature_comparison",
    "promotional_offer":  "promotional_offer",
    "spec_sheet":         "spec_sheet",
    "other":              "general",
}


def detect_document_type(document_type_field: str, filename: str) -> str:
    """
    Determine the extraction prompt to use.

    Priority order:
      1. Upload field (user explicitly chose) — most reliable
      2. Filename keyword matching
      3. Default → "brochure" (most common car dealership document)

    Note: Content-based detection (DOC_TYPE_CONTENT_SIGNALS) is used separately
    in auto_detect_from_content() after extraction, to update the doc type in DB.
    """
    # Priority 1: explicit upload field
    mapped = UPLOAD_FIELD_MAP.get(document_type_field)
    if mapped and mapped != "general":
        logger.info(f"Document type from upload field: '{document_type_field}' → prompt='{mapped}'")
        return mapped

    # Priority 2: filename keywords
    fname = filename.lower().replace("-", " ").replace("_", " ")
    for doc_type, keywords in FILENAME_TYPE_MAP.items():
        if any(kw in fname for kw in keywords):
            logger.info(f"Document type from filename '{filename}': '{doc_type}'")
            return doc_type

    # Priority 3: default
    logger.info(f"Document type defaulting to 'brochure' for '{filename}'")
    return "brochure"


def auto_detect_from_content(text: str) -> Optional[str]:
    """
    Score extracted text against DOC_TYPE_CONTENT_SIGNALS to confirm / correct
    the document type detected from filename alone.

    Returns the best-matched type, or None if no clear winner (score < 3).
    Called after extraction, before chunking.
    """
    text_lower = text.lower()
    scores: Dict[str, int] = {}
    for doc_type, signals in DOC_TYPE_CONTENT_SIGNALS.items():
        score = sum(1 for sig in signals if sig in text_lower)
        if score > 0:
            scores[doc_type] = score

    if not scores:
        return None

    best_type = max(scores, key=lambda k: scores[k])
    best_score = scores[best_type]

    if best_score >= 3:
        logger.info(
            f"Content-based detection: '{best_type}' (score={best_score}) "
            f"| all scores: {scores}"
        )
        return best_type

    logger.info(
        f"Content detection inconclusive (best score={best_score} < 3) "
        f"— keeping filename-based type"
    )
    return None


# ══════════════════════════════════════════════════════════════════════════════
# BOILERPLATE FILTER
# ══════════════════════════════════════════════════════════════════════════════

# Legal disclaimers, trademark notices, and footer text that add zero value
# to the ElevenLabs AI agent's knowledge base and confuse its answers.
BOILERPLATE_PATTERNS = [
    r"reserves the right to change without (prior )?notice",
    r"images?\s+(are\s+)?for (reference|illustration) purpose",
    r"may not be part of standard equipment",
    r"colours?\s+may vary.{0,80}printing",
    r"trademarks?\s+of\s+(apple|amazon|google|android)",
    r"apple\s+carplay\s+is\s+(a\s+trademark|available\s+in)",
    r"android\s+auto.{0,100}trademark",
    r"get the android auto app on google play",
    r"amazon,?\s+alexa,?\s+and all related marks",
    r"maruti\s+suzuki\s+india\s+(ltd|limited),?\s+plot\s+no",
    r"toll\s+free\s+no[:\s]+18",
    r"for\s+more\s+details,?\s+contact\s+toll\s+free",
    r"fuel efficiency as certified by test agency",
    r"spare wheel material is steel",
    r"seat belt reminder comes on unless",
    r"if the vehicle.{0,30}speed exceeds",
    r"buzzer does not indicate any malfunction",
    r"variant.{0,30}features.{0,30}refer to (the )?equipment list",
    r"accessories and features shown in the pictures",
    r"registered in the u\.?s\.? and other countries",
    r"^\s*\*{1,4}[^*\n]{0,120}$",                 # footnote lines starting with *
    r"^\s*#[^#\n]{0,120}(rpm|kg|km)\s*$",          # footnote lines starting with #
    r"^\s*\^\^?[^\^\n]{0,120}$",                   # footnote lines starting with ^
    r"page\s+\d+\s+of\s+\d+",
    r"©\s*\d{4}",
    r"all rights reserved",
    r"printed in india",
    r"visit us at www\.",
    r"follow us on (facebook|instagram|twitter|youtube)",
]
_BOILERPLATE_RE = re.compile(
    "|".join(BOILERPLATE_PATTERNS), re.IGNORECASE | re.MULTILINE
)

# These signals mean the chunk has REAL data — don't discard even if boilerplate matched
_REAL_DATA_SIGNALS = [
    "km/l", "km/kg", "kmpl", "bhp", "ps @", "nm @", "cc",
    "₹", "lakh", "ex-showroom",
    "airbag", "sunroof", "smartplay", "android auto", "apple carplay",
    "lxi", "vxi", "zxi", "alpha", "delta", "sigma", "zeta",
    "warranty", "free service",
    "length:", "width:", "wheelbase:", "boot space:",
]


def _is_boilerplate(text: str) -> bool:
    """
    Return True if chunk is pure legal/trademark/footer boilerplate.
    A chunk that ALSO contains real car data is preserved.
    """
    if not _BOILERPLATE_RE.search(text):
        return False
    text_lower = text.lower()
    if any(sig in text_lower for sig in _REAL_DATA_SIGNALS):
        return False
    return True


# ══════════════════════════════════════════════════════════════════════════════
# SECTION KEYWORDS (used by classify_section)
# ══════════════════════════════════════════════════════════════════════════════

SECTION_KEYWORDS: Dict[str, List[str]] = {
    "pricing": [
        "price", "ex-showroom", "on-road", "emi", "finance", "cost",
        "lakh", "₹", "variant price", "booking", "rto", "insurance charges",
        "amount", "rate", "subvention", "zero down",
    ],
    "features": [
        "feature", "infotainment", "touchscreen", "sunroof", "camera",
        "cruise control", "climate control", "wireless", "android auto",
        "apple carplay", "smartplay", "ambient", "charging dock",
        "paddle shifter", "voice assistant", "suzuki connect", "ota update",
        "surround sound", "arkamys", "head up display", "360 view",
    ],
    "safety": [
        "safety", "airbag", "abs", "esp", "ncap", "star rating", "brake",
        "seatbelt", "seat belt", "collision", "isofix", "hill hold", "ebd",
        "tect body", "impact sensing", "reverse camera", "parking sensor",
        "rear view camera", "pre-tensioner",
    ],
    "specifications": [
        "engine", "displacement", "capacity cc", "bhp", "ps @", "torque", "nm @",
        "mileage", "km/l", "km/kg", "kmpl", "transmission", "wheelbase",
        "ground clearance", "boot space", "fuel tank", "kerb weight",
        "length", "width", "height", "seating capacity", "dimensions",
        "tyre size", "suspension", "ventilated disc", "macpherson",
    ],
    "overview": [
        "overview", "introduction", "about", "highlights", "key points",
        "why choose", "all-new", "designed for", "power to play",
        "more power", "redefine", "new chapter", "next generation",
    ],
    "warranty": [
        "warranty", "service", "maintenance", "annual", "years coverage",
        "km coverage", "roadside", "amc", "free service", "service schedule",
        "whichever is earlier", "labour", "spare parts",
    ],
    "offers": [
        "offer", "discount", "cashback", "exchange", "bonus", "festive",
        "promotion", "scheme", "subvention", "emi", "save", "benefit",
        "limited period", "corporate discount", "loyalty",
    ],
    "colors": [
        "color", "colour", "shade", "arctic white", "metallic", "pearl",
        "dual-tone", "dual tone", "splendid silver", "sizzling red",
        "exuberant blue", "magma grey", "brave khaki", "bluish black",
        "roof", "monotone",
    ],
    "insurance": [
        "insurance", "idv", "insured declared value", "premium", "ncb",
        "no claim bonus", "third party", "comprehensive", "zero depreciation",
        "claim", "policy", "insurer", "sum insured",
    ],
    "accessories": [
        "accessory", "accessories", "floor mat", "seat cover", "body cover",
        "alloy wheel", "dash cam", "fitment", "genuine part", "add-on",
    ],
}


# ══════════════════════════════════════════════════════════════════════════════
# EXTRACTION PROMPTS — one per document type
# ══════════════════════════════════════════════════════════════════════════════

_BASE_RULES = """
UNIVERSAL EXTRACTION RULES:
- Preserve ALL numbers exactly: ₹, km/l, km/kg, bhp, PS, Nm, mm, kg, cc, RPM, litres, %
- Keep model/variant names exact: LXi, VXi, ZXi, ZXi+, Alpha, Delta, Sigma, Zeta, S, V, VX, MT, AT, CNG
- Use "## SECTION NAME" for every major section heading
- Tables/grids: output as pipe-separated rows — Header1 | Header2 | Header3
- Tick marks (✓ ✔ ● Yes): write "Yes". Cross marks (✗ — No dash): write "No"
- Cell values with specific text (e.g. "Bi-Halogen", "22.86cm SmartPlay Pro+"): preserve exactly
- SKIP: page borders, backgrounds, watermarks, decorative art, lifestyle photo captions with no data
- SKIP: legal disclaimer paragraphs (starting with "reserves the right", "for reference purpose", "images are")
- SKIP: trademark footnotes (Apple, Google, Amazon, Android boilerplate)
- SKIP: repeated brand headers/footers appearing on every page
- SKIP: page numbers, print dates, website/social media links at footer
- If a page is a full-bleed lifestyle photo with no data text: write [VISUAL PAGE: brief description]
- Return ONLY extracted content. No preamble, explanation, or meta-commentary.
"""

PROMPTS: Dict[str, str] = {

"brochure": """You are extracting a CAR BROCHURE for an AI voice sales agent.
Customers on the phone will ask: "What features does it have?", "Tell me about the engine",
"What safety does it offer?", "What colours are available?", "What is the mileage?"

Extract EVERY section completely. Use ## headers for each section.

## OVERVIEW / HIGHLIGHTS
- Tagline, key selling proposition, introductory paragraph

## TECHNOLOGY & INFOTAINMENT
- Every tech feature with full name and description
- Screen size and system name (e.g. "22.86cm SmartPlay Pro+ with Surround Sense powered by ARKAMYS")
- Connectivity: Android Auto, Apple CarPlay — wired or wireless
- Voice assistant details (e.g. "Hi Suzuki" barge-in feature)
- Suzuki Connect / connected car features listed individually
- Wireless charging, OTA updates, remote operations

## ENGINE & PERFORMANCE
- All fuel variants: Petrol, Diesel, CNG, Hybrid — each with engine name, capacity, power, torque
- Mileage per variant and transmission:
  Format exactly: "LXi MT: 17.80 km/l | ZXi+ MT: 19.89 km/l | CNG: 25.51 km/kg"
- Smart Hybrid details if present (MHEV, ISG, brake energy regeneration)
- S-CNG system details if present (dual ECU, intelligent injection)

## EXTERIOR
- Headlamp type per variant (Bi-Halogen, Dual LED Projector, LED DRLs)
- Alloy wheel design and size (Painted / Precision Cut / Steel)
- Body cladding, skid plates, roof rails, shark fin antenna
- Dual-tone exterior and which variants get it

## INTERIOR
- Seat material and colour theme per variant (Mono Tone / Dual Tone, fabric/leather)
- Steering wheel type (flat-bottom, leather-wrapped, tilt/telescopic)
- AC type (manual / auto climate control) and rear AC vents
- Ambient lighting, overhead console, armrests, storage features
- Paddle shifters, push start, keyless entry

## SAFETY
- Total number of airbags and placement (Front, Side, Curtain)
- Active safety: ESP, ABS, EBD, Hill Hold Assist, ISOFIX, Impact Sensing Door Unlock
- NCAP rating if mentioned
- Parking aids: rear camera, reverse sensors, 360 view camera
- Seatbelt features: pre-tensioners, force limiters, reminders

## COLOURS
For each colour:
  "Colour Name | Dual Tone: Yes/No | Roof Colour (if dual tone) | Available in: [variants]"
Example: "Sizzling Red with Bluish Black Roof | Dual Tone: Yes | Available in: ZXi, ZXi+"

## WARRANTY & SERVICE
- Standard warranty (years and km)
- Extended warranty options
- Free service details if mentioned
""" + _BASE_RULES,


"pricing_sheet": """You are extracting a CAR PRICING SHEET for an AI voice sales agent.
Customers will ask: "What is the price of ZXi?", "What's the on-road price in Hyderabad?",
"What's the EMI for ZXi+ AT?", "What's the booking amount?"

Extract ALL pricing data with maximum precision.

## PRICING TABLE
For every variant: "Variant | Ex-Showroom Price | On-Road Price (if shown)"
Example:
Variant | Ex-Showroom | On-Road
Swift LXi MT | ₹6,49,000 | ₹7,12,500
Swift VXi MT | ₹7,34,000 | ₹8,05,200
Swift ZXi MT | ₹8,69,000 | ₹9,48,800
Swift ZXi+ AT | ₹9,99,000 | ₹10,88,500

Include ALL variants: petrol MT, petrol AT, CNG, diesel/hybrid if present.

## APPLICABLE CHARGES (if shown)
- RTO / Road Tax charges
- Insurance (1-year / 5-year)
- Extended warranty cost
- Accessories package price
- Handling / logistics charges
- Fastag

## EMI & FINANCE OPTIONS
For each scheme: "Variant | Down Payment | Monthly EMI | Tenure | Interest Rate | Bank"
Example: "ZXi+ AT | ₹1,50,000 | ₹18,499/month | 60 months | 8.5% | HDFC Bank"

## SPECIAL SCHEMES (if present)
- 0% interest EMI, subvention schemes
- Corporate pricing, fleet pricing
- Exchange bonus deduction shown in price

## CITY / LOCATION
- City or state these prices apply to (important — prices vary by location)

## VALIDITY
- Price validity date: "Prices valid till: 31 March 2026"
- Revision note if present
""" + _BASE_RULES,


"feature_comparison": """You are extracting a CAR FEATURE COMPARISON TABLE for an AI voice agent.
Customers will ask: "Does the VXi have a sunroof?", "Which variant gets wireless Android Auto?",
"Does LXi have automatic climate control?", "What screen size does ZXi have?"

This document has multi-variant comparison tables. Extract them completely and accurately.

## FEATURE TABLE — [SECTION NAME]
Use one table block per section. Use the EXACT variant names from the document header row.
Feature | LXi MT | VXi MT | ZXi MT | ZXi+ MT | LXi CNG | VXi CNG | ZXi CNG
Electric Sunroof | No | No | Yes | Yes | No | No | Yes
Touch Screen | No | 17.78cm SmartPlay Studio | 17.78cm SmartPlay Pro | 22.86cm SmartPlay Pro+ | No | 17.78cm SmartPlay Studio | 17.78cm SmartPlay Pro

CRITICAL:
- Extract the ACTUAL cell value, not just Yes/No when the cell contains specific text
  e.g. "Bi-Halogen", "Dual LED", "Manual", "Auto Fold", "(Tilt)", "(Tilt & Telescopic)"
- Group features by their section exactly as in the document:
  SAFETY | INFOTAINMENT | COMFORT AND CONVENIENCE | EXTERIORS | INTERIORS | COLOUR VARIANTS

## SPECIFICATION TABLE (if present)
Extract all specs with variant-specific values:
Spec | LXi/VXi | ZXi/ZXi+
Engine | K15C | K15C Smart Hybrid
Transmission | 5MT / 6AT | 5MT / 6AT
Fuel Efficiency | 17.80 km/l (MT) / 19.80 km/l (AT) | 19.89 km/l (MT) / 19.80 km/l (AT)

## COLOR AVAILABILITY (if present)
Colour | LXi | VXi | ZXi | ZXi+
Pearl Arctic White | Yes | Yes | Yes | Yes
Sizzling Red with Bluish Black Roof | No | No | Yes | Yes
""" + _BASE_RULES,


"promotional_offer": """You are extracting a PROMOTIONAL OFFERS DOCUMENT for an AI voice agent.
Customers will ask: "What offers are available?", "How much discount can I get?",
"Is there an exchange bonus?", "Any festive offer?", "What's the total saving?"

Extract every offer with complete details so the agent can quote them precisely.

## CURRENT OFFERS
For each offer type, extract:
- Offer Type: Cash Discount / Exchange Bonus / Corporate Discount / Loyalty Bonus / Accessories Voucher
- Amount: ₹X,XXX or X%
- Applicable Variants: "All variants" or list specific ones
- Eligibility / Conditions
- Stacking: can this be combined with other offers?

Format:
"Cash Discount: ₹30,000 | Applicable: ZXi, ZXi+ | Condition: Available to all retail customers"
"Exchange Bonus: ₹20,000 | Applicable: All variants | Condition: On exchange of any old vehicle"
"Corporate Discount: ₹15,000 | Applicable: All variants | Condition: Valid employee ID required"

## EMI & FINANCE SCHEMES
"0% EMI: ZXi+ AT | 12 months | Down payment: ₹1,50,000 | Bank: HDFC"
"Low EMI: ZXi MT | ₹14,499/month | 60 months | 8.5% p.a. | Bank: SBI"

## TOTAL BENEFIT CALCULATOR (if shown)
"Maximum total saving: ₹75,000 on ZXi+"

## VALIDITY & TERMS
- Valid till: [date]
- Stock subject to availability
- Dealership-specific conditions
- Contact number for offer enquiries
""" + _BASE_RULES,


"spec_sheet": """You are extracting a CAR SPECIFICATION SHEET for an AI voice agent.
Customers will ask: "What is the engine size?", "What's the mileage?",
"How big is the boot?", "What are the dimensions?", "Is it BS6?"

Extract all technical specifications completely and clearly.

## DIMENSIONS
Length: [mm] | Width: [mm] | Height (unladen): [mm] | Wheelbase: [mm]
Ground Clearance: [mm] | Kerb Weight: [kg] | Boot Space: [litres]
Seating Capacity: [number]

## ENGINE — PETROL (if present)
Engine Name: | Type: | Capacity: [cc]
Max Power: [kW / bhp / PS @ RPM]
Max Torque: [Nm @ RPM]
Emission Norm: BS VI
Smart Hybrid: Yes/No (with type: MHEV / ISG)

## ENGINE — CNG (if present)
Engine Name: | Capacity: [cc]
Max Power (Petrol): [kW / PS @ RPM] | Max Power (CNG): [kW / PS @ RPM]
Max Torque (Petrol): [Nm @ RPM] | Max Torque (CNG): [Nm @ RPM]
CNG Tank Capacity: [litres water equivalent]
Petrol Tank Capacity: [litres]

## ENGINE — DIESEL (if present)
[same format as petrol]

## FUEL EFFICIENCY (per variant and transmission)
LXi MT: [km/l] | VXi MT: [km/l] | ZXi MT: [km/l] | ZXi+ MT: [km/l]
VXi AT: [km/l] | ZXi AT: [km/l] | ZXi+ AT: [km/l]
CNG variants: [km/kg]

## TRANSMISSION
Available types and which variants get which:
"5MT: LXi, VXi, ZXi, ZXi+ | 6AT: VXi, ZXi, ZXi+"

## BRAKES
Front: [Ventilated Disc / Solid Disc / Drum]
Rear: [Drum / Disc]

## SUSPENSION
Front: [MacPherson Strut / Double Wishbone]
Rear: [Torsion Beam / Multi-Link]

## TYRES
Size: [215/60 R16] | Type: [Tubeless Radial]
Spare: [Steel / Alloy — size]

## FUEL TANK
Petrol: [litres] | CNG: [litres water equivalent] | Diesel: [litres]
""" + _BASE_RULES,


"warranty": """You are extracting a WARRANTY / SERVICE SCHEDULE document for an AI voice agent.
Customers will ask: "How long is the warranty?", "What's covered?",
"How many free services?", "What does AMC include?"

## WARRANTY COVERAGE
Standard Warranty: [X years or X,XX,XXX km — whichever is earlier]
Extended Warranty Options:
  - [X years / X,XX,XXX km] — Cost: ₹[amount] (if shown)
  - [X years / X,XX,XXX km] — Cost: ₹[amount]
Powertrain Warranty (if separate): [details]
Battery Warranty (hybrid/EV): [details]

## FREE SERVICE SCHEDULE
1st Free Service: [1,000 km or 1 month — whichever is earlier]
2nd Free Service: [10,000 km or 12 months]
3rd Free Service: [20,000 km or 24 months]
[continue for all scheduled services]

## PAID SERVICE SCHEDULE (if shown)
After free services, list paid service intervals:
"Every 10,000 km or 12 months — Labour: ₹X,XXX | Parts estimate: ₹X,XXX"

## ANNUAL MAINTENANCE CONTRACT (AMC)
Packages available: [list]
What's included: [oil change, filters, labour, etc.]
What's excluded: [tyres, batteries, accidental damage, etc.]
Price: ₹[amount] for [duration]

## ROADSIDE ASSISTANCE
Coverage: [24x7 / Pan-India / X km radius]
Services included: [towing, flat tyre, fuel delivery, battery jump-start]
Contact: [number or app name]

## WHAT'S NOT COVERED (key exclusions)
[Wear and tear items, consumables, accidental damage, misuse, etc.]
""" + _BASE_RULES,


"insurance": """You are extracting a CAR INSURANCE DOCUMENT for an AI voice agent.
Customers will ask: "What insurance coverage do I get?", "What's the IDV?",
"How do I make a claim?", "What is zero depreciation?"

## POLICY DETAILS
Policy Number: | Insurer: | Policy Type: [Comprehensive / Third Party]
Policy Period: | Vehicle: | Registration Number:

## COVERAGE
IDV (Insured Declared Value): ₹[amount]
Own Damage Coverage: ₹[amount]
Third Party Liability: [As per Motor Vehicles Act]
Personal Accident Cover (Owner-Driver): ₹[amount]

## ADD-ON COVERS (if listed)
- Zero Depreciation: Yes/No — ₹[premium]
- Engine Protect: Yes/No — ₹[premium]
- Roadside Assistance: Yes/No
- Return to Invoice: Yes/No
- Consumables Cover: Yes/No
- NCB Protect: Yes/No

## PREMIUM BREAKUP (if shown)
Own Damage Premium: ₹[amount]
Third Party Premium: ₹[amount]
Add-on Premium: ₹[amount]
Total Premium (excl. GST): ₹[amount]
GST (18%): ₹[amount]
Total Premium (incl. GST): ₹[amount]

## NO CLAIM BONUS (NCB)
Current NCB: [0% / 20% / 25% / 35% / 45% / 50%]
NCB Slab: 1st renewal: 20% | 2nd: 25% | 3rd: 35% | 4th: 45% | 5th+: 50%

## CLAIM PROCESS
1. [Step-by-step claim instructions]
Claim helpline: [number]
Cashless garages: [network details]
""" + _BASE_RULES,


"accessory": """You are extracting a CAR ACCESSORIES CATALOGUE for an AI voice agent.
Customers will ask: "What accessories are available?", "How much does the dash cam cost?",
"Do you have seat covers?", "What alloy wheels can I get?"

## ACCESSORIES LIST
For each accessory:
"Accessory Name | Part Number (if shown) | Price: ₹[amount] | Fitment: [Dealer / DIY]"

Group by category:

## EXTERIOR ACCESSORIES
[Body cover, mud flaps, side cladding, chrome garnish, roof rails, etc.]

## INTERIOR ACCESSORIES
[Floor mats, seat covers, steering cover, gear knob, organiser, etc.]

## TECHNOLOGY ACCESSORIES
[Dash cam, GPS tracker, parking sensors, rear camera, etc.]

## PERFORMANCE / PROTECTION
[Tyre pressure monitor, battery maintainer, anti-rust treatment, etc.]

## ALLOY WHEELS (if listed)
"[Size] [Design Name] | Painted / Machined | Price per wheel: ₹[amount] | Set of 4: ₹[amount]"

## FITMENT PACKAGES (if shown)
[Bundled package names, contents, and total price]
""" + _BASE_RULES,


"general": """You are extracting a car dealership document for an AI voice sales agent.
The document may contain any combination of: brochure content, pricing, features, specs, or offers.

Extract ALL content a customer on the phone might ask about.

Identify and use ## SECTION HEADERS for:
- Pricing / Variant Prices
- Features & Technology
- Engine & Performance / Mileage
- Safety Features
- Interior & Exterior Design
- Colours & Variants
- Offers & Discounts
- Warranty & Service
- Specifications & Dimensions
- Any other relevant section

For tables: use pipe-separated format — Col1 | Col2 | Col3
For mileage callout boxes: "Variant: X.XX km/l"
For feature tick/cross grids: "Feature | Variant1: Yes | Variant2: No"
Preserve all numbers exactly as shown: ₹, km/l, km/kg, bhp, PS, Nm, mm, cc
""" + _BASE_RULES,

}


# ══════════════════════════════════════════════════════════════════════════════
# OPENAI CLIENT
# ══════════════════════════════════════════════════════════════════════════════

def _get_openai_client():
    if not OPENAI_API_KEY:
        raise RuntimeError(
            "OPENAI_API_KEY not set.\n"
            "Get a key at https://platform.openai.com/api-keys\n"
            "Then add to .env: OPENAI_API_KEY=sk-..."
        )
    try:
        from openai import OpenAI  # type: ignore[import]
        return OpenAI(api_key=OPENAI_API_KEY)
    except ImportError:
        raise RuntimeError(
            "openai SDK not installed.\n"
            "Run: pip install openai --break-system-packages"
        )


# ══════════════════════════════════════════════════════════════════════════════
# RATE LIMITING
# ══════════════════════════════════════════════════════════════════════════════

def _rate_limit_openai() -> None:
    global _last_openai_call_time
    now = time.monotonic()
    elapsed = now - _last_openai_call_time
    if elapsed < OPENAI_MIN_INTERVAL_SECONDS:
        time.sleep(OPENAI_MIN_INTERVAL_SECONDS - elapsed)
    _last_openai_call_time = time.monotonic()


# ══════════════════════════════════════════════════════════════════════════════
# CORE EXTRACTION — image bytes → OpenAI Vision
# ══════════════════════════════════════════════════════════════════════════════

def _extract_from_image_bytes(
    image_bytes: bytes,
    mime_type: str,
    client,
    prompt: str,
    page_label: str = "page",
    max_retries: int = 3,
) -> str:
    """
    Send image bytes + prompt to OpenAI Vision. Returns extracted text or "".

    Features:
    - MD5 cache: same image never sent twice in same process
    - Proactive rate limiting before every call
    - Exponential backoff on 429 rate limit errors
    - Quota exhausted guard: flips _openai_quota_exhausted flag, no more retries
    """
    cache_key = hashlib.md5(image_bytes).hexdigest()
    if cache_key in _page_cache:
        logger.info(f"Cache hit: {page_label}")
        return _page_cache[cache_key]

    b64_image = base64.b64encode(image_bytes).decode("utf-8")
    data_url = f"data:{mime_type};base64,{b64_image}"

    for attempt in range(1, max_retries + 1):
        try:
            _rate_limit_openai()

            response = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": data_url, "detail": "high"}},
                    ],
                }],
                max_tokens=2000,
                temperature=0,
            )

            text = (response.choices[0].message.content or "").strip()
            logger.info(f"OpenAI ✓ {page_label} → {len(text)} chars")
            _page_cache[cache_key] = text
            return text

        except Exception as e:
            err_str = str(e)

            if any(x in err_str for x in ["401", "authentication", "api_key", "Incorrect API key"]):
                logger.error(f"OpenAI auth failed — check OPENAI_API_KEY: {e}")
                return ""

            if "insufficient_quota" in err_str or "exceeded your current quota" in err_str:
                global _openai_quota_exhausted
                _openai_quota_exhausted = True
                logger.error("OpenAI quota exhausted — switching all pages to PyMuPDF fallback")
                return ""

            if ("429" in err_str or "rate_limit" in err_str.lower()) and attempt < max_retries:
                retry_match = re.search(r"retry.{0,20}?(\d+)\s*s", err_str, re.IGNORECASE)
                wait = min(int(retry_match.group(1)) + 2 if retry_match else 10 * attempt, 65)
                logger.warning(f"Rate limit on {page_label} (attempt {attempt}/{max_retries}). Waiting {wait}s")
                time.sleep(wait)
                continue

            if attempt < max_retries:
                time.sleep(3 * attempt)
            else:
                logger.warning(f"OpenAI failed on {page_label} after {max_retries} attempts: {e}")
                return ""

    return ""


# ══════════════════════════════════════════════════════════════════════════════
# PDF EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_path: str, doc_type: str = "brochure") -> str:
    """
    Extract text from a PDF.

    PYMUPDF_ONLY mode: pure PyMuPDF text extraction (fast, free, no API).
    Production mode: OpenAI Vision per page, PyMuPDF fallback on failure.

    Why Vision for car brochures:
      - Multi-column layouts → PyMuPDF linearises incorrectly
      - Feature comparison tables → PyMuPDF merges cells, loses Yes/No structure
      - Mileage callout boxes → PyMuPDF misses them on styled pages
      - Image-only pages (e.g. colour swatches) → PyMuPDF returns empty
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise RuntimeError("PyMuPDF not installed. Run: pip install pymupdf --break-system-packages")

    pdf_doc = fitz.open(file_path)
    total_pages = len(pdf_doc)  # type: ignore[arg-type]
    process_pages = min(total_pages, MAX_PAGES)
    page_texts: List[str] = []
    stats = {"openai": 0, "pymupdf": 0, "visual_only": 0}

    # ── PyMuPDF-only mode (dev/testing) ───────────────────────────────────────
    if PYMUPDF_ONLY:
        logger.info(f"PDF (PyMuPDF-only): '{file_path}' | {process_pages}/{total_pages} pages | type={doc_type}")
        for page_num in range(process_pages):
            page = pdf_doc[page_num]
            text = page.get_text("text").strip()  # type: ignore[attr-defined]
            if text and len(text) >= MIN_TEXT_THRESHOLD:
                page_texts.append(f"[Page {page_num + 1}]\n{text}")
                stats["pymupdf"] += 1
            else:
                page_texts.append(f"[Page {page_num + 1}] [VISUAL — no extractable text]")
                stats["visual_only"] += 1
        pdf_doc.close()
        logger.info(f"✓ PyMuPDF complete | text={stats['pymupdf']} pages | visual={stats['visual_only']} pages")

    # ── OpenAI Vision primary, PyMuPDF fallback ───────────────────────────────
    else:
        client = _get_openai_client()
        prompt = PROMPTS.get(doc_type, PROMPTS["general"])
        logger.info(f"PDF (OpenAI Vision): '{file_path}' | {process_pages}/{total_pages} pages | type={doc_type}")

        for page_num in range(process_pages):
            page = pdf_doc[page_num]
            label = f"page {page_num + 1}/{process_pages}"

            mat = fitz.Matrix(PAGE_DPI / 72, PAGE_DPI / 72)  # type: ignore[attr-defined]
            pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)  # type: ignore[attr-defined]
            img_bytes = pix.tobytes("png")

            openai_text = "" if _openai_quota_exhausted else _extract_from_image_bytes(
                img_bytes, "image/png", client, prompt, label
            )

            if openai_text and len(openai_text.strip()) > 30:
                page_texts.append(f"[Page {page_num + 1}]\n{openai_text}")
                stats["openai"] += 1
            else:
                pymupdf_text = page.get_text("text").strip()  # type: ignore[attr-defined]
                if pymupdf_text and len(pymupdf_text) >= MIN_TEXT_THRESHOLD:
                    page_texts.append(f"[Page {page_num + 1}] [pymupdf-fallback]\n{pymupdf_text}")
                    stats["pymupdf"] += 1
                else:
                    page_texts.append(f"[Page {page_num + 1}] [VISUAL — no extractable text]")
                    stats["visual_only"] += 1

        pdf_doc.close()
        logger.info(
            f"✓ PDF complete | OpenAI={stats['openai']} | PyMuPDF={stats['pymupdf']} | Visual={stats['visual_only']}"
        )

    if total_pages > MAX_PAGES:
        page_texts.append(f"[Note: {total_pages} pages total. Only first {MAX_PAGES} processed.]")

    return "\n\n".join(page_texts)


# ══════════════════════════════════════════════════════════════════════════════
# DOCX EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_docx(file_path: str, doc_type: str = "brochure") -> str:
    """
    Extract DOCX:
      - Text paragraphs and tables via python-docx (no API call)
      - Embedded images → OpenAI Vision (skipped in PYMUPDF_ONLY mode)
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx --break-system-packages")

    client = None if PYMUPDF_ONLY else _get_openai_client()
    prompt = PROMPTS.get(doc_type, PROMPTS["general"])
    doc = DocxDocument(file_path)
    parts: List[str] = []

    # Text paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = (para.style.name or "") if para.style else ""
        if style_name.startswith("Heading"):
            level = style_name.replace("Heading", "").strip()
            prefix = "##" if level in ("1", "2", "") else "###"
            parts.append(f"\n{prefix} {text}\n")
        else:
            parts.append(text)

    # Tables → pipe-separated rows
    for table in doc.tables:
        rows: List[str] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # De-duplicate merged cells (python-docx repeats merged cell content)
            deduped: List[str] = []
            prev = None
            for c in cells:
                if c != prev:
                    deduped.append(c)
                prev = c
            if any(deduped):
                rows.append(" | ".join(deduped))
        if rows:
            parts.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")

    # Embedded images → OpenAI Vision
    if not PYMUPDF_ONLY and client is not None:
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    mime = "image/png" if image_data[:4] == b'\x89PNG' else "image/jpeg"
                    image_count += 1
                    img_text = _extract_from_image_bytes(
                        image_data, mime, client, prompt, f"embedded image {image_count}"
                    )
                    if img_text and "[VISUAL PAGE" not in img_text and len(img_text) > 20:
                        parts.append(f"\n[From embedded image {image_count}]\n{img_text}")
                except Exception as e:
                    logger.warning(f"Could not extract DOCX image: {e}")
    else:
        logger.info("DOCX: skipping embedded image extraction (PYMUPDF_ONLY mode)")

    return "\n\n".join(parts)


# ══════════════════════════════════════════════════════════════════════════════
# IMAGE EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_image(file_path: str, doc_type: str = "brochure") -> str:
    """
    Send an uploaded image directly to OpenAI Vision.
    In PYMUPDF_ONLY mode, logs a warning and returns a placeholder.
    """
    if PYMUPDF_ONLY:
        logger.warning(
            f"Image uploaded in PYMUPDF_ONLY mode: '{file_path}'. "
            "Set PYMUPDF_ONLY=false and add OPENAI_API_KEY to extract text from images."
        )
        return (
            "[Image document uploaded. Text extraction requires Vision API. "
            "Set OPENAI_API_KEY and PYMUPDF_ONLY=false in .env to enable.]"
        )

    client = _get_openai_client()
    prompt = PROMPTS.get(doc_type, PROMPTS["general"])

    with open(file_path, "rb") as f:
        image_bytes = f.read()

    ext = os.path.splitext(file_path)[1].lower()
    mime_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png", ".webp": "image/webp"}
    mime_type = mime_map.get(ext, "image/jpeg")

    text = _extract_from_image_bytes(image_bytes, mime_type, client, prompt, "uploaded image")
    if not text:
        raise ValueError("OpenAI Vision could not extract any text from this image")
    return text


# ══════════════════════════════════════════════════════════════════════════════
# TEXT CLEANING
# ══════════════════════════════════════════════════════════════════════════════

def clean_extracted_text(raw_text: str) -> str:
    """Normalise extracted text before chunking."""
    if not raw_text:
        return ""

    # Remove pipeline page markers
    text = re.sub(r'\[Page \d+(/\d+)?\]\s*(\[.*?\])?\n?', '', raw_text)

    # Remove TABLE markers (keep content)
    text = re.sub(r'\[/?TABLE\]\n?', '', text)

    # Collapse 3+ blank lines → 2
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Remove standalone page numbers on their own line
    text = re.sub(r'^\s*\d{1,3}\s*$', '', text, flags=re.MULTILINE)

    # Remove "Page X of Y"
    text = re.sub(r'Page\s+\d+\s+of\s+\d+', '', text, flags=re.IGNORECASE)

    # Collapse multiple spaces (preserve newlines and tabs)
    text = re.sub(r'[ \t]{2,}', ' ', text)

    # Strip non-printable chars (keep unicode, tabs, newlines)
    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]', '', text)

    return text.strip()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION CLASSIFICATION
# ══════════════════════════════════════════════════════════════════════════════

def classify_section(text: str) -> str:
    """
    Classify a chunk into a section type.
    
    Priority:
    1. ## section header in chunk text (strong signal)
    2. Keyword scoring across all section types
    3. Default → "general"
    """
    # Check for ## header — most reliable signal
    header_match = re.search(r'##\s+([\w][\w\s&+/\-]*)', text)
    if header_match:
        header = header_match.group(1).lower()
        header_map = {
            "pricing":        ["price", "variant price", "cost", "rate", "emi", "finance", "booking"],
            "specifications": ["spec", "engine", "dimension", "performance", "fuel", "technical", "transmission"],
            "safety":         ["safety", "airbag", "esp", "ncap"],
            "features":       ["feature", "technology", "connect", "infotainment", "interior", "exterior",
                               "comfort", "convenience", "playful", "smartplay"],
            "offers":         ["offer", "discount", "promo", "scheme", "emi", "benefit", "saving", "cashback"],
            "warranty":       ["warranty", "service", "amc", "maintenance"],
            "colors":         ["color", "colour", "shade"],
            "overview":       ["overview", "about", "highlight", "introduction", "power to play"],
            "insurance":      ["insurance", "policy", "idv", "premium", "ncb"],
            "accessories":    ["accessor", "fitment", "add-on", "parts"],
        }
        for section_type, keywords in header_map.items():
            if any(kw in header for kw in keywords):
                return section_type

    # Keyword scoring fallback
    text_lower = text.lower()
    scores: Dict[str, int] = {}
    for section, keywords in SECTION_KEYWORDS.items():
        score = sum(1 for kw in keywords if kw in text_lower)
        if score > 0:
            scores[section] = score

    return max(scores, key=lambda k: scores[k]) if scores else "general"


# ══════════════════════════════════════════════════════════════════════════════
# CHUNKING
# ══════════════════════════════════════════════════════════════════════════════

def split_into_chunks(text: str) -> List[Tuple[int, str]]:
    """
    Split text into meaningful chunks for the knowledge base.

    Strategy (in order):
    1. Split on ## section markers — each section becomes a chunk candidate
    2. If a section is too large (> 1.5x CHUNK_SIZE), split by paragraphs
    3. If a paragraph is too large, split by sentences

    Filters applied:
    - MIN_CHUNK_CHARS: drops fragment/orphan chunks
    - _is_boilerplate: drops legal disclaimers and trademark noise
    """
    chunks: List[str] = []

    if "## " in text:
        sections = re.split(r'\n(?=## )', text)
        for section in sections:
            section = section.strip()
            if not section:
                continue
            if len(section) <= CHUNK_SIZE * 1.5:
                chunks.append(section)
            else:
                chunks.extend(_split_by_paragraphs(section))
    else:
        chunks = _split_by_paragraphs(text)

    return [
        (i, c.strip())
        for i, c in enumerate(chunks)
        if len(c.strip()) >= MIN_CHUNK_CHARS and not _is_boilerplate(c)
    ]


def _split_by_paragraphs(text: str) -> List[str]:
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    chunks: List[str] = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) < CHUNK_SIZE:
            current = (current + "\n\n" + para).strip()
        else:
            if current:
                chunks.append(current)
            if len(para) > CHUNK_SIZE:
                sub = _split_by_sentences(para)
                chunks.extend(sub[:-1])
                current = sub[-1] if sub else ""
            else:
                current = para
    if current:
        chunks.append(current)
    return chunks


def _split_by_sentences(text: str) -> List[str]:
    sentences = re.split(r'(?<=[.!?])\s+', text)
    chunks: List[str] = []
    current = ""
    for s in sentences:
        if len(current) + len(s) < CHUNK_SIZE:
            current = (current + " " + s).strip()
        else:
            if current:
                chunks.append(current)
            current = s
    if current:
        chunks.append(current)
    return chunks


# ══════════════════════════════════════════════════════════════════════════════
# MAIN PROCESSING FUNCTION
# ══════════════════════════════════════════════════════════════════════════════

def process_document(document_id: str, db: Session) -> dict:
    """
    Full document processing pipeline:

    1. Load document record from DB
    2. Detect document type → select appropriate extraction prompt
    3. Extract text (PDF / DOCX / Image)
    4. Content-based type re-detection (may upgrade "general" → specific type)
    5. Clean extracted text
    6. Split into chunks, filter boilerplate
    7. Classify each chunk's section type
    8. Bulk save chunks to DB
    9. Update document: processing_status=completed, processed_at, chunk_count

    All errors are caught, DB is rolled back, and document is marked failed.
    """
    doc = db.query(Document).filter(Document.document_id == document_id).first()
    if not doc:
        return {"error": "Document not found", "document_id": str(document_id)}

    # Mark as in-progress
    doc.processing_status = "processing"   # type: ignore[assignment]
    doc.processing_error  = None           # type: ignore[assignment]
    doc.updated_at        = datetime.utcnow()  # type: ignore[assignment]
    db.commit()

    extraction_method = "unknown"

    try:
        file_path = str(doc.file_path)
        file_type = str(doc.file_type)
        filename  = str(doc.filename or "")

        # Step 1: Detect document type from upload field + filename
        doc_type = detect_document_type(
            str(doc.document_type or "brochure"),
            filename
        )

        logger.info(
            f"Processing: '{filename}' | "
            f"file_type={file_type} | detected_doc_type={doc_type} | "
            f"PYMUPDF_ONLY={PYMUPDF_ONLY}"
        )

        # Step 2: Extract text
        if file_type == "pdf":
            raw_text = extract_text_from_pdf(file_path, doc_type)
            extraction_method = (
                "pymupdf-only" if PYMUPDF_ONLY
                else f"openai-vision/{OPENAI_MODEL} + pymupdf-fallback"
            )
        elif file_type == "docx":
            raw_text = extract_text_from_docx(file_path, doc_type)
            extraction_method = (
                "python-docx (text only)" if PYMUPDF_ONLY
                else f"python-docx + openai-vision/{OPENAI_MODEL}"
            )
        elif file_type == "image":
            raw_text = extract_text_from_image(file_path, doc_type)
            extraction_method = (
                "none (PYMUPDF_ONLY mode — Vision API needed for images)" if PYMUPDF_ONLY
                else f"openai-vision/{OPENAI_MODEL}"
            )
        else:
            raise ValueError(f"Unsupported file type: '{file_type}'")

        # Step 3: Content-based type re-detection (after extraction)
        # This can correct "general" → specific type when filename gave no signal
        if doc_type in ("brochure", "general"):
            content_type = auto_detect_from_content(raw_text)
            if content_type and content_type != doc_type:
                logger.info(
                    f"Content detection upgraded doc_type: '{doc_type}' → '{content_type}'"
                )
                doc_type = content_type

        # Step 4: Clean
        clean_text = clean_extracted_text(raw_text)
        if not clean_text or len(clean_text) < MIN_CHUNK_CHARS:
            raise ValueError(
                f"No readable text extracted from '{filename}'. "
                "If this is an image-heavy PDF, ensure OPENAI_API_KEY is set and "
                "PYMUPDF_ONLY=false in your .env file."
            )

        # Step 5: Store full extracted text on document
        doc.extracted_text = clean_text  # type: ignore[assignment]

        # Step 6: Chunk
        chunk_tuples = split_into_chunks(clean_text)
        if not chunk_tuples:
            raise ValueError(
                f"Document '{filename}' could not be split into usable chunks. "
                "The content may be entirely boilerplate or too short."
            )

        # Step 7: Delete old chunks (handles reprocessing)
        deleted_count = db.query(DocumentChunk).filter(
            DocumentChunk.document_id == document_id
        ).delete()
        if deleted_count > 0:
            logger.info(f"Deleted {deleted_count} old chunks for reprocessing")

        # Step 8: Save new chunks
        for idx, chunk_text in chunk_tuples:
            chunk = DocumentChunk(
                chunk_id      = uuid.uuid4(),
                document_id   = doc.document_id,
                dealership_id = doc.dealership_id,
                car_model_id  = doc.car_model_id,
                chunk_index   = idx,
                chunk_text    = chunk_text,
                section_type  = classify_section(chunk_text),
                char_count    = len(chunk_text),
            )
            db.add(chunk)

        # Step 9: Finalise document record
        now = datetime.utcnow()
        doc.processing_status = "completed"        # type: ignore[assignment]
        doc.chunk_count       = len(chunk_tuples)  # type: ignore[assignment]
        doc.processed_at      = now                # type: ignore[assignment]  ← was always null before
        doc.updated_at        = now                # type: ignore[assignment]
        db.commit()

        logger.info(
            f"✅ '{filename}': {len(chunk_tuples)} chunks | "
            f"type={doc_type} | method={extraction_method}"
        )

        return {
            "document_id":       str(document_id),
            "processing_status": "completed",
            "chunk_count":       len(chunk_tuples),
            "doc_type_detected": doc_type,
            "extraction_method": extraction_method,
            "message": (
                f"Extracted {len(chunk_tuples)} chunks from '{filename}' "
                f"[{doc_type}] via {extraction_method}"
            ),
        }

    except Exception as e:
        db.rollback()
        err_msg = str(e)
        doc.processing_status = "failed"           # type: ignore[assignment]
        doc.processing_error  = err_msg            # type: ignore[assignment]
        doc.updated_at        = datetime.utcnow()  # type: ignore[assignment]
        db.commit()
        logger.error(f"❌ '{doc.filename}' processing failed: {err_msg}")
        return {
            "document_id":       str(document_id),
            "processing_status": "failed",
            "chunk_count":       0,
            "extraction_method": extraction_method,
            "message":           f"Processing failed: {err_msg}",
        }